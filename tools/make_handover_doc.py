"""Generate the CFR27 sim handover doc (.docx) for Teams. Short, plain, no waffle."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)

d = Document()
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(10.5)


def H(txt):
    h = d.add_heading(txt, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def P(txt, bold=False, size=10.5, color=None):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(txt)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def B(txt):
    p = d.add_paragraph(txt, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def table(rows, widths=None):
    t = d.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for par in cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)
                    if i == 0:
                        r.bold = True
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    d.add_paragraph()
    return t


t = d.add_heading("CFR27 Gear Ratio Study", level=0)
for r in t.runs:
    r.font.color.rgb = NAVY
P("What final drive should we run? Built from real car data, not guesses.",
  size=12, color=GREY)
P("Repo: github.com/Alr0611/CFR-gear-ratio-optimizations", size=9, color=GREY)

H("The answer")
P("Lower ratios help efficiency and pack charge. Higher ratios help acceleration. "
  "4.61 is the middle, which is why CFR24 picked it.")
P("Dropping to 4.2 gets us about 7 more points of energy spent in the motor's efficient "
  "range, plus a bit more charge left at the end of endurance. It costs about 0.19 s over "
  "75 m. That cost is small because our launch is limited by tyre grip, not motor torque - "
  "the gearing isn't what's holding us back off the line.")
P("I'd go 4.4-4.5. Most of the efficiency, almost none of the accel.", bold=True)
P("If we want 4.2, a 21T/38T gearbox on our current chain gets us 4.17 and the gears come "
  "out stronger than what we run now (bigger pinion, less tooth load).")

H("How to run it")
P("Open MATLAB in the repo folder, add it to the path with subfolders, then:")
table([
    ["Command", "What you get"],
    ["gear_ratio_optimization", "The main study - efficiency and pack charge per ratio, plus 4 figures"],
    ["accel_model", "Accel - 0-75 m, 0-100 kph, and how much tyre weight matters"],
    ["open_system('accel_sim')", "Same accel model in Simulink"],
    ["fatigue_spectrum / accel_fatigue", "Driveline torque load spectra"],
    ["verify_math", "24 checks that recompute everything from scratch"],
], widths=[2.3, 4.5])
P("Every constant is in params_cfr26.m. Change it there and everything updates. "
  "Figures land in output/.", size=9, color=GREY)

H("Where the numbers come from")
table([
    ["Input", "Source"],
    ["Efficiency and operating points", "Comp Jun 20 endurance"],
    ["Pack charge, battery model", "July 11 test (the only complete run - comp DNF'd)"],
    ["Accel check", "Comp Jun 19 launches, real 0-75 m = 4.40 s"],
    ["Torque load spectra", "Comp Jun 20 + Jun 19"],
    ["Downforce and balance", "Ford wind tunnel"],
    ["Mass, CG, weight split", "Tilt test, with driver"],
    ["Motor", "EMRAX 208 datasheet"],
    ["Cells", "HPPC test + ESF (BAK 45D)"],
], widths=[2.4, 4.4])

H("How much to trust it")
B("Battery model tracks the real pack to about 5 mV per cell over 80 minutes.")
B("Motor efficiency is built from datasheet physics and lands on the datasheet's own 96% "
  "peak on its own. It was never tuned to look right.")
B("Accel model says 4.72 s where the real launch was 4.40 s. It runs slightly slow, never "
  "optimistic. Simulink and MATLAB agree to 0.02 s.")
B("Starting charge from rest voltage (94.7%) matches what the BMS said (94.3%).")
B("verify_math.m runs 24 checks against the source documents. All pass.")

H("Things we found along the way")
B("Accel is what fatigues the driveline, not endurance. Endurance peaks at 132 Nm and "
  "basically never goes above 120. Accel peaks at 152 Nm and spends 23% of its time above "
  "120. The old fatigue sheet only went to 140 Nm so it couldn't even see this.")
B("Our CFD over-predicts downforce by about 65%. The Ford tunnel measured L/D 1.61; CFD "
  "implied 2.68. If CFD downforce is being used anywhere else, it's too high.")
B("No regen, confirmed from brake data. About 25% of our traction energy goes out as heat "
  "through the brakes. Gearing can't recover any of that.")
B("Drag barely matters at our speeds - top speed is redline limited, not drag limited.")
B("The 159 Nm peak was partly a sensor glitch. Real peak is 152 Nm.")

H("What NOT to trust")
B("Tyre longitudinal grip is derived, not measured - Calspan never tested our tyre that "
  "way. So accel times are a range, roughly plus/minus 15%, not a promise.")
B("Comp endurance DNF'd on a hot cell, so pack charge comes from July 11 instead.")
B("Drivetrain efficiency (82.3%), rolling resistance and wheel inertia are estimates.")
B("Everything assumes July 11 pace is competition pace.")
P("None of these change which ratio wins. They move the absolute numbers.", size=9, color=GREY)

H("Gear strength - not finished, don't quote it")
P("There's a Gear Check workbook in the repo. It reproduces the CFR24 Driveline Tool's "
  "stress math, but I can't close it out yet, so treat it as a work in progress.")
P("The problem: the CFR24 tool doesn't add up. It says 20 degree pressure angle and its I "
  "factor agrees, but its J factor (0.325) is a 25 degree value - the 20 degree chart reads "
  "0.245 at 15 teeth. On its own stated geometry, our bending stress should be 470 MPa, not "
  "the 354 it reports.")
P("The gears are probably fine anyway. A 15 tooth pinion at 20 degrees would be undercut "
  "unless it's profile shifted, and shifting it adds root material and pushes J up to about "
  "where CFR24 has it. So 0.325 might be right for a reason nobody wrote down. Can't tell "
  "from the spreadsheet.")
P("To close this out we need two numbers off the gear drawing: pressure angle, and the "
  "profile shift coefficient. Also what steel the gears are and the heat treat - right now "
  "the allowables in there are Baja's, not ours.", bold=True)
P("None of this affects the ratio recommendation. 21T/38T sits above the undercut limit "
  "either way.", size=9, color=GREY)

H("Open items")
table([
    ["Item", "Who"],
    ["Gear drawing: pressure angle + profile shift", "whoever spec'd the gears"],
    ["Gear material + heat treat", "same"],
    ["Cell effective capacity (coulomb count vs BMS disagree ~6 points)", "elec"],
    ["CFD vs tunnel downforce correlation", "aero"],
    ["Pick the ratio", "us"],
], widths=[4.6, 2.2])

P("Questions - Aboud", size=9, color=GREY)

out = r"c:\Users\Aboud\Downloads\CFR27 Gear Ratio Study.docx"
d.save(out)
print("Saved:", os.path.basename(out))
