"""START HERE doc for the Teams copy. Short. Plain. Rewrite it in your own words."""
import os, sys, subprocess
from docx import Document
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
RED  = RGBColor(0x9C, 0x00, 0x06)

SHA = sys.argv[1] if len(sys.argv) > 1 else "unknown"
WHEN = sys.argv[2] if len(sys.argv) > 2 else "unknown"

d = Document()
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(10.5)


def H(t):
    h = d.add_heading(t, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY


def P(t, bold=False, size=10.5, color=None):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def B(t):
    p = d.add_paragraph(t, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)


t = d.add_heading("Start here", level=0)
for r in t.runs:
    r.font.color.rgb = NAVY
P("CFR27 gear ratio study - what's in this folder and what to open first.", size=12, color=GREY)

H("If you only open one thing")
P("CFR27 Gear Ratio Study.docx - the whole thing in about a page. Read that first.")

H("What's in here")
B("CFR27 Gear Ratio Study.docx - the report. Start here.")
B("MATLAB sim - all the code. Open MATLAB in this folder, add to path with subfolders, "
  "run gear_ratio_optimization. Figures land in output/.")
B("Excel tools - Fatigue Load Cases is finished and rebuilt from real telemetry. "
  "Gear Check is NOT finished, see below.")
B("Legacy - CFR24 - the old Driveline Tool, kept for reference. Read the warning below "
  "before you trust anything in it.")

H("This is a copy, not the original")
P("The real version lives on GitHub: github.com/Alr0611/CFR-gear-ratio-optimizations")
P("Snapshot of commit %s, taken %s." % (SHA, WHEN), size=9, color=GREY)
P("If you edit files in this Teams folder, your changes will NOT make it back to the repo, "
  "and the next time someone re-uploads a snapshot they'll be wiped. Want something "
  "changed? Message me and I'll do it properly so it sticks.", bold=True)

H("Two things not to quote")
P("Gear Check.xlsx is a work in progress.", bold=True)
P("It reproduces the CFR24 Driveline Tool's stress math, but I can't close it out. The "
  "CFR24 tool doesn't add up: it states 20 degree pressure angle and its I factor agrees, "
  "but its J factor is a 25 degree value. On its own stated geometry our bending stress "
  "should be 470 MPa, not the 354 it reports. The gears are probably fine anyway - a 15 "
  "tooth pinion at 20 degrees has to be profile shifted or it'd be undercut, and shifting "
  "raises J to about where CFR24 has it. But that's a guess, not a fact.")
P("Also, the material allowables in there are Baja's steel (SCM415), not ours. So the "
  "factors of safety are indicative only. Comparing two gearsets is fine - the allowable "
  "cancels out. Quoting an absolute FOS is not.")
P("To close it out we need two things off the gear drawing: pressure angle and profile "
  "shift coefficient. Plus what steel our gears actually are and the heat treat.", bold=True)
P("None of this changes the ratio recommendation.", size=9, color=GREY)

H("The CFR24 Driveline Tool - read this before using it")
P("It looks authoritative and it cost us a lot of time. Known issues:")
B("Its J factor doesn't match its own stated pressure angle (see above).")
B("It has no allowables and no factors of safety at all - it computes stresses and stops.")
B("Its fatigue S-N columns were never finished (marked 0% complete in the tool itself).")
P("It's here for reference and history, not as a source of truth.", size=9, color=GREY)

H("Questions")
P("Ask Aboud. If I've graduated, the report has a 'what not to trust' section and every "
  "constant in the sim is in params_cfr26.m with a comment saying where it came from. "
  "Start there.")

out = r"c:\Users\Aboud\Downloads\Gear Ratio Study\00 START HERE.docx"
d.save(out)
print("Saved:", out)
